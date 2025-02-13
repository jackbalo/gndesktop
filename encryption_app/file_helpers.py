import re
import unicodedata
import math
from docx import Document
import os
import uuid
import threading
import shutil
from flask import current_app

UPLOAD_PICS_FOLDER = os.path.join(os.path.dirname(__file__), 'static/profile_pics')
ALLOWED_PIC_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def allowed_pic_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_PIC_EXTENSIONS

UPLOAD_DOCS_FOLDER = os.path.join(os.path.dirname(__file__), 'static/Documents')
ALLOWED_DOC_EXTENSIONS = {'docx'}

def allowed_doc_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_DOC_EXTENSIONS

def create_document(content):
    try:
        filename = f"{uuid.uuid4()}.docx"
        file_path = os.path.join(UPLOAD_DOCS_FOLDER, filename)
        
        # Ensure the directory exists
        os.makedirs(UPLOAD_DOCS_FOLDER, exist_ok=True)
        
        doc = Document()
        doc.add_paragraph(content)
        doc.save(file_path)
        return filename
    except Exception as e:
        current_app.logger.error(f"Error creating document: {e}")
        raise

def duplicate_document(file_path):
    try:
        filename = os.path.basename(file_path)
        duplicate_filename = f"{os.path.splitext(filename)[0]}_encrypt{os.path.splitext(filename)[1]}"
        duplicate_file_path = os.path.join(UPLOAD_DOCS_FOLDER, duplicate_filename)
        shutil.copy(file_path, duplicate_file_path)
        return duplicate_file_path
    except Exception as e:
        current_app.logger.error(f"Error duplicating document: {e}")
        raise

def normalize_text(text):
    text = unicodedata.normalize('NFC', text)
    # Define punctuation replacements
    replacements = {
        "–":"-", "—":"-", "(.)":".", "“":'"', "”":'"', "‘":'"', "’":"'", "…":"...", "•":"*", "—":"-", ':': ':', ';': ';', '.': '.', ',': ',', '"': '"', "'": "'", '(': '(', ')': ')', '/': '/', '-': '-', ' ': ' ', '\n': '\n'
    }
    

    for symbol, replacement in replacements.items():
        text = text.replace(symbol, replacement)

    return text


def copy_section(file_path):
    try:
        doc = Document(file_path)
        
        # initialize flag to indicate when section starts
        section_found = False
        section_text = ""

        # predefined list of keywords
        keywords = ["REST'D", "RESTRICTED", "SEC", "SECRET", "TOP SECRET", "TOPSECRET", "TOPSEC", "CONF", "CONFIDENTIAL", "GR", "GR:"]
        normalized_keywords = [normalize_text(keyword) for keyword in keywords]

        # iterate through paragraphs and tables
        for para in doc.paragraphs:
            para_text = normalize_text(para.text)
            if section_found:
                section_text += para_text + "\n"
                if "//" in para_text:
                    break
            elif any(para_text.startswith(keyword) for keyword in normalized_keywords):
                section_found = True
                section_text += para_text + "\n"

        for table in doc.tables:
            for row in table.rows:
                row_text = ",".join(normalize_text(cell.text) for cell in row.cells)
                if section_found:
                    section_text += row_text + "\n"
                    if "//" in row_text:
                        break
                elif any(normalize_text(cell.text).startswith(keyword) for keyword in normalized_keywords):
                    section_found = True
                    section_text += row_text + "\n"
        
        # Ensure section_text is not empty
        if not section_text.strip():
            raise ValueError("No matching section found in the document.")
        
        # Remove the last line if it matches the pattern
        lines = section_text.strip().split("\n")
        if re.match(r"\d{6}Z/[A-Z]{3}/\d{2}", lines[-1].strip()) or re.match(r"\d{6}Z [A-Z]{3} \d{2}", lines[-1].strip()):
            lines = lines[:-1]
        
        # Remove the group count line if it matches the pattern
        if lines[0].strip().startswith("GR:"):
            lines = lines[1:]
        
        return "\n".join(lines).strip()
    except Exception as e:
        current_app.logger.error(f"Error copying section: {e}")
        raise

def replace_section(file_path, edited_text):
    try:
        doc = Document(file_path)
        
        # initialize flag to indicate when section starts
        section_found = False

        # predefined list of keywords
        keywords = ["REST'D", "RESTRICTED", "SEC", "SECRET", "TOP SECRET", "TOPSECRET", "TOPSEC", "CONF", "CONFIDENTIAL", "GR", "GR:"]  # Keywords to identify the section
        normalized_keywords = [normalize_text(keyword) for keyword in keywords]

        # iterate through paragraphs and tables
        for para in doc.paragraphs:
            para_text = normalize_text(para.text)
            if section_found:
                if "//" in para_text:
                    para.text = edited_text
                    section_found = False
                    break
                para.text = ""
            elif any(para_text.startswith(keyword) for keyword in normalized_keywords):
                section_found = True
                para.text = ""

        if section_found:
            for table in doc.tables:
                for row in table.rows:
                    row_text = ",".join(normalize_text(cell.text) for cell in row.cells)
                    if section_found:
                        if "//" in row_text:
                            for cell in row.cells:
                                cell.text = edited_text
                            section_found = False
                            break
                        for cell in row.cells:
                            cell.text = ""
                    elif any(normalize_text(cell.text).startswith(keyword) for keyword in normalized_keywords for cell in row.cells):
                        section_found = True
                        for cell in row.cells:
                            cell.text = ""

        if section_found:
            doc.add_paragraph(edited_text)

        doc.save(file_path)
    except Exception as e:
        current_app.logger.error(f"Error replacing section: {e}")
        raise

def delete_file(file_path, delay):
    def delete():
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            current_app.logger.error(f"Error deleting file: {e}")
    timer = threading.Timer(delay, delete)
    timer.start()